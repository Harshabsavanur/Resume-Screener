import io
import re
import base64
import math
import hashlib
import sqlite3
import os
import json
import requests
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from flask import Flask, request, jsonify
from flask_cors import CORS
from dotenv import load_dotenv

# Load environment variables
load_dotenv()
OPENROUTER_API_KEY = os.environ.get("OPENROUTER_API_KEY")
OPENROUTER_MODEL = os.environ.get("OPENROUTER_MODEL", "openrouter/free")

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

app = Flask(__name__, static_folder='.', static_url_path='')
# Enable CORS so the local frontend can communicate with this API
CORS(app)

@app.route('/')
def index_route():
    return app.send_static_file('index.html')

DB_PATH = 'screening.db'

def get_db_connection():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def hash_password(password):
    """Hashes a password using SHA-256 with a fixed salt for simplicity."""
    salted = 'nexusnlp_salt_' + password
    return hashlib.sha256(salted.encode('utf-8')).hexdigest()

def init_db():
    with get_db_connection() as conn:
        conn.execute("PRAGMA foreign_keys = ON;")
        conn.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            full_name TEXT NOT NULL,
            email TEXT NOT NULL UNIQUE,
            password_hash TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );
        """)
        conn.execute("""
        CREATE TABLE IF NOT EXISTS jobs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            skills TEXT,
            min_cgpa REAL,
            min_experience INTEGER,
            description TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );
        """)
        conn.execute("""
        CREATE TABLE IF NOT EXISTS candidates (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            job_id INTEGER NOT NULL,
            name TEXT NOT NULL,
            filename TEXT NOT NULL,
            overall_score REAL,
            status TEXT,
            skills_match_score REAL,
            project_relevance_score REAL,
            experience_score REAL,
            cgpa_score REAL,
            extracted_cgpa TEXT,
            extracted_experience INTEGER,
            matched_skills TEXT,
            missing_skills TEXT,
            notes TEXT DEFAULT '',
            summary TEXT DEFAULT '',
            shortlist_status TEXT DEFAULT 'Pending',
            email TEXT DEFAULT '',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (job_id) REFERENCES jobs (id) ON DELETE CASCADE
        );
        """)
        
        # Ensure 'summary' column exists for existing databases
        try:
            conn.execute("ALTER TABLE candidates ADD COLUMN summary TEXT DEFAULT '';")
        except sqlite3.OperationalError:
            pass

        # Ensure 'shortlist_status' column exists for existing databases
        try:
            conn.execute("ALTER TABLE candidates ADD COLUMN shortlist_status TEXT DEFAULT 'Pending';")
        except sqlite3.OperationalError:
            pass

        # Ensure 'email' column exists for existing databases
        try:
            conn.execute("ALTER TABLE candidates ADD COLUMN email TEXT DEFAULT '';")
        except sqlite3.OperationalError:
            pass

        # Seed a default admin user if no users exist
        existing = conn.execute("SELECT COUNT(*) FROM users").fetchone()[0]
        if existing == 0:
            conn.execute(
                "INSERT INTO users (full_name, email, password_hash) VALUES (?, ?, ?)",
                ('Admin', 'admin@nexusnlp.com', hash_password('admin123'))
            )
        conn.commit()

init_db()


# Standard English stopwords to filter out for basic NLP processing
STOPWORDS = {
    'a', 'about', 'above', 'after', 'again', 'against', 'all', 'am', 'an', 'and', 'any', 'are', 'aren', "aren't",
    'as', 'at', 'be', 'because', 'been', 'before', 'being', 'below', 'between', 'both', 'but', 'by', 'can', 'cannot',
    'could', 'couldn', "couldn't", 'did', 'didn', "didn't", 'do', 'does', 'doesn', "doesn't", 'doing', 'don', "don't",
    'down', 'during', 'each', 'few', 'for', 'from', 'further', 'had', 'hadn', "hadn't", 'has', 'hasn', "hasn't",
    'have', 'haven', "haven't", 'having', 'he', "he'd", "he'll", "he's", 'her', 'here', "here's", 'hers', 'herself',
    'him', 'himself', 'his', 'how', "how's", 'i', "i'd", "i'll", "i'm", "i've", 'if', 'in', 'into', 'is', 'isn', "isn't",
    'it', "it's", 'its', 'itself', 'let', "let's", 'me', 'more', 'most', 'mustn', "mustn't", 'my', 'myself', 'no',
    'nor', 'not', 'of', 'off', 'on', 'once', 'only', 'or', 'other', 'ought', 'our', 'ours', 'ourselves', 'out',
    'over', 'own', 'same', 'shan', "shan't", 'she', "she'd", "she'll", "she's", 'should', 'shouldn', "shouldn't",
    'so', 'some', 'such', 'than', 'that', "that's", 'the', 'their', 'theirs', 'them', 'themselves', 'then', 'there',
    "there's", 'these', 'they', "they'd", "they'll", "they're", "they've", 'this', 'those', 'through', 'to', 'too',
    'under', 'until', 'up', 'very', 'was', 'wasn', "wasn't", 'we', "we'd", "we'll", "we're", "we've", 'were', 'weren',
    "weren't", 'what', "what's", 'when', "when's", 'where', "where's", 'which', 'while', 'who', "who's", 'whom',
    'why', "why's", 'with', 'won', "won't", 'would', 'wouldn', "wouldn't", 'you', "you'd", "you'll", "you're", "you've",
    'your', 'yours', 'yourself', 'yourselves'
}

# --- Parsing Helpers ---

def extract_text_from_pdf(file_bytes):
    if PdfReader is None:
        raise ImportError("pypdf is not installed on the system.")
    
    pdf_file = io.BytesIO(file_bytes)
    reader = PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        extracted = page.extract_text()
        if extracted:
            text += extracted + "\n"
    return text

def extract_text_from_docx(file_bytes):
    if Document is None:
        raise ImportError("python-docx is not installed on the system.")
    
    docx_file = io.BytesIO(file_bytes)
    doc = Document(docx_file)
    text_parts = []
    # Read paragraphs
    for para in doc.paragraphs:
        if para.text:
            text_parts.append(para.text)
    # Read table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    text_parts.append(cell.text)
    return "\n".join(text_parts)

def extract_text_from_txt(file_bytes):
    # Try common encodings
    for encoding in ['utf-8', 'latin-1', 'cp1252']:
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("Failed to decode text file with standard encodings.")

# --- NLP Processing & Similarity Helpers ---

def clean_text(text):
    """Lowercases, removes punctuation/special characters, and tokenizes while filtering stopwords."""
    if not text:
        return []
    
    # Lowercase
    text_lower = text.lower()
    
    # Replace special characters and punctuation with spaces, but preserve numbers/letters
    cleaned = re.sub(r'[^a-z0-9\s\+\#\.]', ' ', text_lower)
    
    # Split on whitespace
    tokens = cleaned.split()
    
    # Filter stopwords and short terms (except special symbols like C, C++, C#)
    filtered = []
    for token in tokens:
        # Strip trailing dot if it's not part of a common acronym/decimal
        if token.endswith('.') and not re.match(r'^\d+\.\d+$', token):
            token = token[:-1]
            
        if token and token not in STOPWORDS:
            filtered.append(token)
            
    return filtered

def compute_cosine_similarity(text1, text2):
    """Computes cosine similarity based on simple term frequencies (bag of words)."""
    words1 = clean_text(text1)
    words2 = clean_text(text2)
    
    if not words1 or not words2:
        return 0.0
        
    tf1 = {}
    for w in words1:
        tf1[w] = tf1.get(w, 0) + 1
        
    tf2 = {}
    for w in words2:
        tf2[w] = tf2.get(w, 0) + 1
        
    unique_words = set(tf1.keys()).union(set(tf2.keys()))
    
    dot_product = sum(tf1.get(w, 0) * tf2.get(w, 0) for w in unique_words)
    magnitude1 = math.sqrt(sum(count ** 2 for count in tf1.values()))
    magnitude2 = math.sqrt(sum(count ** 2 for count in tf2.values()))
    
    if magnitude1 == 0 or magnitude2 == 0:
        return 0.0
        
    return dot_product / (magnitude1 * magnitude2)

# --- Feature Extraction Helpers ---

def extract_skills(resume_text, target_skills_list):
    """Matches target skills in a case-insensitive manner against the resume."""
    text_lower = resume_text.lower()
    matched = []
    missing = []
    
    for skill in target_skills_list:
        skill_clean = skill.strip().lower()
        if not skill_clean:
            continue
        
        # Escape skill for safe regex insertion
        escaped_skill = re.escape(skill_clean)
        
        # For skills with punctuation like C++, C#, .NET, we do substring check.
        # Otherwise, check with word boundaries to avoid false positives (e.g. "Go" matching "Good").
        if not re.match(r'^\w', skill_clean) or not re.match(r'.*\w$', skill_clean) or '+' in skill_clean or '#' in skill_clean or '.' in skill_clean:
            found = skill_clean in text_lower
        else:
            # Word boundary regex matching
            pattern = r'\b' + escaped_skill + r'\b'
            found = re.search(pattern, text_lower) is not None
            
        if found:
            # Keep original case from user input
            matched.append(skill.strip())
        else:
            missing.append(skill.strip())
            
    score = (len(matched) / len(target_skills_list) * 100) if target_skills_list else 100.0
    return matched, missing, score

def extract_cgpa(text):
    """Extracts CGPA out of 10 or conversion from percentage from the resume."""
    text_lower = text.lower()
    
    # Patterns to match CGPA / GPA patterns
    patterns = [
        r'\b(?:cgpa|gpa|pointer|c\.g\.p\.a|g\.p\.a)[:\s-]*([0-9]\.[0-9]{1,2})\b',
        r'\b(?:cgpa|gpa|pointer)[:\s-]*([0-9]{1,2}(?:\.[0-9]{1,2})?)\s*/\s*10\b',
        r'\b([0-9]\.[0-9]{1,2})\s*/\s*10\b',
        r'\b([0-9]\.[0-9]{1,2})\s*(?:cgpa|gpa|pointer)\b',
        r'\b(?:cgpa|gpa|pointer|grade)[:\s-]*([0-9]\.[0-9]{1,2})\b'
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, text_lower)
        if matches:
            try:
                val = float(matches[0])
                if 0.0 <= val <= 10.0:
                    return val
            except ValueError:
                continue
                
    # Fallback to percentage parsing (e.g. 85%) and converting it to 10-point scale
    pct_pattern = r'\b(\d{2}(?:\.\d{1,2})?)\s*%'
    pct_matches = re.findall(pct_pattern, text_lower)
    if pct_matches:
        try:
            pct = float(pct_matches[0])
            if 35.0 <= pct <= 100.0:
                # Convert to 10-point scale: 85% becomes 8.5
                return round(pct / 10.0, 2)
        except ValueError:
            pass
            
    # Search for general numbers between 5.0 and 10.0 that look like a CGPA
    float_pattern = r'\b([5-9]\.[0-9]{1,2})\b'
    float_matches = re.findall(float_pattern, text_lower)
    if float_matches:
        try:
            val = float(float_matches[0])
            return val
        except ValueError:
            pass
            
    return None

def extract_experience(text):
    """Extracts candidate years of experience based on keywords."""
    text_lower = text.lower()
    
    patterns = [
        r'\b(\d+)\+?\s*(?:year|yr)s?\s*(?:of\s*)?experience\b',
        r'\b(\d+)\+?\s*(?:year|yr)s?\s*exp\b',
        r'\bexperience[:\s-]*(\d+)\+?\s*(?:year|yr)s?\b',
        r'\b(\d+)\+?\s*(?:year|yr)s?\s*in\b',
        r'\bwork\s*experience[:\s-]*(\d+)\+?\s*(?:year|yr)s?\b'
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, text_lower)
        if matches:
            try:
                return int(matches[0])
            except ValueError:
                continue
                
    # Scan text window around word "experience" to look for numbers
    match = re.search(r'experience', text_lower)
    if match:
        start = max(0, match.start() - 50)
        end = min(len(text_lower), match.end() + 50)
        window = text_lower[start:end]
        num_matches = re.findall(r'\b(\d+)\b', window)
        for num in num_matches:
            val = int(num)
            if 1 <= val <= 25:
                return val
                
    return 0

def extract_name(text, filename):
    """Extracts candidate's name from resume content or falls back to filename."""
    if not text:
        return re.sub(r'\.[^.]+$', '', filename).replace('_', ' ').replace('-', ' ').title()
        
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    for line in lines[:4]:
        # Skip if contains email, links, or common headers
        if '@' in line or 'http' in line or re.search(r'\b(experience|education|skills|profile|resume|cv|email|phone|contact|tel|mobile)\b', line, re.I):
            continue
        words = line.split()
        if 2 <= len(words) <= 4 and all(w.isalpha() or w.endswith('.') for w in words):
            return line
            
    return re.sub(r'\.[^.]+$', '', filename).replace('_', ' ').replace('-', ' ').title()


def extract_email(text):
    """Extracts candidate's email address from resume text using regular expressions."""
    if not text:
        return ""
    # Standard email regex pattern
    pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    matches = re.findall(pattern, text)
    if matches:
        return matches[0].strip()
    return ""


# --- NLP LLM & Screening Helpers ---

def analyze_resume_with_llm(resume_content, is_text, mime_type, job_title, skills_raw, min_cgpa, min_exp, job_desc):
    if not OPENROUTER_API_KEY:
        raise ValueError("OPENROUTER_API_KEY not configured in .env file.")
        
    url = "https://openrouter.ai/api/v1/chat/completions"
    
    if is_text:
        prompt = f"""You are an expert ATS (Applicant Tracking System) assistant. Analyze the following candidate's resume text against the job requirements.

Job Requirements:
Job Title: {job_title}
Required Skills: {skills_raw}
Minimum CGPA: {min_cgpa}
Minimum Experience: {min_exp} years
Job Description: {job_desc}

Candidate Resume Text:
{resume_content}

Analyze the resume and return a JSON object with the following fields:
- "name": String. Extracted name of the candidate. If not found, use a fallback from filename if provided.
- "email": String. Extracted email address of the candidate, or null if not found.
- "extracted_cgpa": Number (float out of 10.0 scale) or null if not found. If in percentage (e.g., 85%), convert it to a 10-point scale (8.5). If out of 4.0 scale (e.g., 3.6/4), convert it to a 10-point scale (9.0).
- "extracted_experience": Number (integer). Extracted years of experience. If not specified or entry-level, use 0.
- "matched_skills": List of strings. Skills from the Required Skills list that are present in the resume. Match semantically (e.g. "ReactJS" matches "React").
- "missing_skills": List of strings. Skills from the Required Skills list that are NOT present in the resume.
- "summary": String. A brief summary of the candidate's profile, qualifications, and fit for the role in exactly 30 words or less.
- "skills_score": Number (0-100). Be liberal and use a curved evaluation (e.g., matching 2 out of 4 skills should give 70+).
- "relevance_score": Number (0-100). Semantic relevance of candidate's projects/experience to the job description. Be liberal: if the projects are related in domain or show potential, score it high (75+).
- "experience_score": Number (0-100). Be liberal: score candidate high (80+) if candidate's experience is close to min_experience or shows strong skills despite a lower year count. If candidate_experience >= min_experience, use 100.
- "cgpa_score": Number (0-100). Be liberal: if candidate CGPA is not found, default to 85. If it is slightly below min_cgpa, score it 80+. If >= min_cgpa, use 100.
- "overall_score": Number (0-100). Weighted score: 40% skills_score + 30% relevance_score + 15% experience_score + 15% cgpa_score.

Response MUST be a valid JSON object ONLY, adhering to the schema. Do not enclose in markdown blocks.
"""
        messages = [
            {
                "role": "user",
                "content": prompt
            }
        ]
    else:
        prompt = f"""You are an expert ATS (Applicant Tracking System) assistant. Analyze the attached candidate resume (provided as an image or PDF document) against the job requirements.

Job Requirements:
Job Title: {job_title}
Required Skills: {skills_raw}
Minimum CGPA: {min_cgpa}
Minimum Experience: {min_exp} years
Job Description: {job_desc}

Analyze the resume and return a JSON object with the following fields:
- "name": String. Extracted name of the candidate. If not found, use a fallback from filename if provided.
- "email": String. Extracted email address of the candidate, or null if not found.
- "extracted_cgpa": Number (float out of 10.0 scale) or null if not found. If in percentage (e.g., 85%), convert it to a 10-point scale (8.5). If out of 4.0 scale (e.g., 3.6/4), convert it to a 10-point scale (9.0).
- "extracted_experience": Number (integer). Extracted years of experience. If not specified or entry-level, use 0.
- "matched_skills": List of strings. Skills from the Required Skills list that are present in the resume. Match semantically (e.g. "ReactJS" matches "React").
- "missing_skills": List of strings. Skills from the Required Skills list that are NOT present in the resume.
- "summary": String. A brief summary of the candidate's profile, qualifications, and fit for the role in exactly 30 words or less.
- "skills_score": Number (0-100). Be liberal and use a curved evaluation (e.g., matching 2 out of 4 skills should give 70+).
- "relevance_score": Number (0-100). Semantic relevance of candidate's projects/experience to the job description. Be liberal: if the projects are related in domain or show potential, score it high (75+).
- "experience_score": Number (0-100). Be liberal: score candidate high (80+) if candidate's experience is close to min_experience or shows strong skills despite a lower year count. If candidate_experience >= min_experience, use 100.
- "cgpa_score": Number (0-100). Be liberal: if candidate CGPA is not found, default to 85. If it is slightly below min_cgpa, score it 80+. If >= min_cgpa, use 100.
- "overall_score": Number (0-100). Weighted score: 40% skills_score + 30% relevance_score + 15% experience_score + 15% cgpa_score.

Response MUST be a valid JSON object ONLY, adhering to the schema. Do not enclose in markdown blocks.
"""
        messages = [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": prompt
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:{mime_type};base64,{resume_content}"
                        }
                    }
                ]
            }
        ]

    payload = {
        "model": OPENROUTER_MODEL,
        "messages": messages,
        "response_format": {
            "type": "json_object"
        },
        "max_tokens": 300
    }
    
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "HTTP-Referer": "http://localhost:5000",
        "X-Title": "NexusNLP Resume Screener"
    }
    
    models_to_try = [
        OPENROUTER_MODEL,
        "meta-llama/llama-3.3-70b-instruct:free",
        "meta-llama/llama-3.2-3b-instruct:free",
        "google/gemma-2-9b-it:free",
        "qwen/qwen3-coder:free",
        "nousresearch/hermes-3-llama-3.1-405b:free",
    ]
    seen = set()
    models_to_try = [x for x in models_to_try if x and not (x in seen or seen.add(x))]

    last_error = ""
    for model in models_to_try:
        payload["model"] = model
        try:
            response = requests.post(url, headers=headers, json=payload, timeout=8)
            if response.status_code != 200:
                last_error = f"Model {model} returned status {response.status_code}: {response.text}"
                print(last_error)
                continue
                
            res_data = response.json()
            content_text = res_data['choices'][0]['message']['content']
            if not content_text:
                last_error = f"Model {model} returned empty content"
                print(last_error)
                continue
                
            content_text = content_text.strip()
            # Remove potential markdown block markers if models don't respect prompt exactly
            if content_text.startswith("```"):
                lines = content_text.splitlines()
                if lines[0].startswith("```json"):
                    content_text = "\n".join(lines[1:-1])
                elif lines[0].startswith("```"):
                    content_text = "\n".join(lines[1:-1])
            parsed_json = json.loads(content_text)
            return parsed_json
        except Exception as e:
            last_error = f"Model {model} failed: {str(e)}"
            print(last_error)
            continue
            
    raise Exception(f"All attempted OpenRouter models failed. Last error: {last_error}")



def screen_resume(file_bytes, filename, job_title, skills_raw, min_cgpa, min_exp, job_desc):
    # Target skills parsing
    target_skills = [s.strip() for s in skills_raw.split(',') if s.strip()] if skills_raw else []
    
    # 1. Determine Mime Type and File Extension
    ext = filename.split('.')[-1].lower()
    is_image = ext in ['png', 'jpg', 'jpeg']
    
    mime_type = None
    if ext == 'pdf':
        mime_type = 'application/pdf'
    elif ext == 'png':
        mime_type = 'image/png'
    elif ext in ['jpg', 'jpeg']:
        mime_type = 'image/jpeg'
        
    resume_text = ""
    text_extraction_success = False
    
    # Try text extraction if it's not a direct image
    if not is_image:
        try:
            if ext == 'pdf':
                resume_text = extract_text_from_pdf(file_bytes)
                text_extraction_success = True
            elif ext == 'docx':
                resume_text = extract_text_from_docx(file_bytes)
                text_extraction_success = True
            elif ext == 'txt':
                resume_text = extract_text_from_txt(file_bytes)
                text_extraction_success = True
        except Exception as e:
            print(f"Text extraction failed for {filename}: {str(e)}")
            
    # Check if we should use multimodal screening (scanned PDF or direct image)
    is_scanned_pdf = (ext == 'pdf' and (not resume_text or len(resume_text.strip()) < 150))
    is_multimodal = is_image or is_scanned_pdf
    
    if is_multimodal:
        if OPENROUTER_API_KEY:
            try:
                # Base64 encode the file bytes
                encoded_data = base64.b64encode(file_bytes).decode('utf-8')
                
                llm_result = analyze_resume_with_llm(
                    resume_content=encoded_data,
                    is_text=False,
                    mime_type=mime_type,
                    job_title=job_title,
                    skills_raw=skills_raw,
                    min_cgpa=min_cgpa,
                    min_exp=min_exp,
                    job_desc=job_desc
                )
                
                overall_score = llm_result.get('overall_score', 0.0)
                skills_score = llm_result.get('skills_score', 0.0)
                relevance_score = llm_result.get('relevance_score', 0.0)
                exp_score = llm_result.get('experience_score', 0.0)
                cgpa_score = llm_result.get('cgpa_score', 0.0)
                
                extracted_cgpa = llm_result.get('extracted_cgpa')
                extracted_exp = llm_result.get('extracted_experience', 0)
                matched_skills = llm_result.get('matched_skills', [])
                missing_skills = llm_result.get('missing_skills', [])
                summary = llm_result.get('summary', '')
                candidate_name = llm_result.get('name') or extract_name("", filename)
                
                if overall_score >= 72:
                    status = 'Excellent Match'
                elif overall_score >= 60:
                    status = 'Good Match'
                elif overall_score >= 45:
                    status = 'Average Match'
                else:
                    status = 'Low Match'
                    
                return {
                    'overall_score': round(overall_score, 1),
                    'status': status,
                    'candidate_name': candidate_name,
                    'skills_score': round(skills_score, 1),
                    'relevance_score': round(relevance_score, 1),
                    'exp_score': round(exp_score, 1),
                    'cgpa_score': round(cgpa_score, 1),
                    'extracted_cgpa': str(extracted_cgpa if extracted_cgpa is not None else "Not Found"),
                    'extracted_exp': extracted_exp,
                    'matched_skills_str': ",".join(matched_skills),
                    'missing_skills_str': ",".join(missing_skills),
                    'summary': summary,
                    'email': llm_result.get('email') or ''
                }
            except Exception as e:
                print(f"OpenRouter multimodal screening failed for {filename}: {str(e)}")
                # Fail gracefully if LLM fails and it's a scanned PDF/image
                return {
                    'overall_score': 0.0,
                    'status': 'Screening Error',
                    'candidate_name': extract_name("", filename),
                    'skills_score': 0.0,
                    'relevance_score': 0.0,
                    'exp_score': 0.0,
                    'cgpa_score': 0.0,
                    'extracted_cgpa': 'Not Found',
                    'extracted_exp': 0,
                    'matched_skills_str': '',
                    'missing_skills_str': skills_raw,
                    'summary': f"Error parsing image/scanned PDF: {str(e)}. Please check your backend logs.",
                    'email': ''
                }
        else:
            # Fallback when key is missing: explain that key is required for scanned PDFs/images
            return {
                'overall_score': 0.0,
                'status': 'Missing API Key',
                'candidate_name': extract_name("", filename),
                'skills_score': 0.0,
                'relevance_score': 0.0,
                'exp_score': 0.0,
                'cgpa_score': 0.0,
                'extracted_cgpa': 'N/A',
                'extracted_exp': 0,
                'matched_skills_str': '',
                'missing_skills_str': skills_raw,
                'summary': "Multimodal screening for scanned PDFs or images requires a valid OPENROUTER_API_KEY in the backend .env configuration.",
                'email': ''
            }
            
    # For standard text-based PDF, DOCX, and TXT:
    if OPENROUTER_API_KEY:
        try:
            llm_result = analyze_resume_with_llm(
                resume_content=resume_text,
                is_text=True,
                mime_type=None,
                job_title=job_title,
                skills_raw=skills_raw,
                min_cgpa=min_cgpa,
                min_exp=min_exp,
                job_desc=job_desc
            )
            overall_score = llm_result.get('overall_score', 0.0)
            skills_score = llm_result.get('skills_score', 0.0)
            relevance_score = llm_result.get('relevance_score', 0.0)
            exp_score = llm_result.get('experience_score', 0.0)
            cgpa_score = llm_result.get('cgpa_score', 0.0)
            
            extracted_cgpa = llm_result.get('extracted_cgpa')
            extracted_exp = llm_result.get('extracted_experience', 0)
            matched_skills = llm_result.get('matched_skills', [])
            missing_skills = llm_result.get('missing_skills', [])
            summary = llm_result.get('summary', '')
            candidate_name = llm_result.get('name') or extract_name(resume_text, filename)
            
            if overall_score >= 72:
                status = 'Excellent Match'
            elif overall_score >= 60:
                status = 'Good Match'
            elif overall_score >= 45:
                status = 'Average Match'
            else:
                status = 'Low Match'
                
            return {
                'overall_score': round(overall_score, 1),
                'status': status,
                'candidate_name': candidate_name,
                'skills_score': round(skills_score, 1),
                'relevance_score': round(relevance_score, 1),
                'exp_score': round(exp_score, 1),
                'cgpa_score': round(cgpa_score, 1),
                'extracted_cgpa': str(extracted_cgpa if extracted_cgpa is not None else "Not Found"),
                'extracted_exp': extracted_exp,
                'matched_skills_str': ",".join(matched_skills),
                'missing_skills_str': ",".join(missing_skills),
                'summary': summary,
                'email': llm_result.get('email') or ''
            }
        except Exception as e:
            print(f"OpenRouter API screening failed, falling back to regex: {str(e)}")
            
    # Fallback / Original Regex and Cosine Similarity Code
    matched_skills, missing_skills, raw_skills_score = extract_skills(resume_text, target_skills)
    
    # Liberal curved skills score (curved to be more forgiving)
    if target_skills:
        skills_score = round(((len(matched_skills) / len(target_skills)) ** 0.6) * 100.0, 1)
    else:
        skills_score = 100.0
        
    comparison_text = job_desc if (job_desc and job_desc.strip()) else f"{job_title} {skills_raw}"
    raw_similarity = compute_cosine_similarity(resume_text, comparison_text) * 100
    
    # Liberal relevance score calculation
    relevance_score = raw_similarity * 2.5  # amplify low cosine scores
    if len(matched_skills) > 0:
        # Guarantee a baseline score based on skill match
        skills_ratio = len(matched_skills) / max(len(target_skills), 1)
        relevance_score = max(relevance_score, 50.0 + (skills_ratio * 35.0))
    else:
        relevance_score = max(relevance_score, 40.0)
    relevance_score = min(relevance_score, 100.0)
    
    extracted_cgpa = extract_cgpa(resume_text)
    if min_cgpa <= 0:
        cgpa_score = 100.0
    elif extracted_cgpa is None:
        cgpa_score = 85.0  # Liberal default if not specified/found
    elif extracted_cgpa >= min_cgpa:
        cgpa_score = 100.0
    else:
        deficit = min_cgpa - extracted_cgpa
        cgpa_score = max(50.0, 100.0 - (deficit * 30.0))  # Forgiving curve
        
    extracted_exp = extract_experience(resume_text)
    if min_exp <= 0:
        exp_score = 100.0
    elif extracted_exp >= min_exp:
        exp_score = 100.0
    else:
        # Forgiving baseline + experience percentage
        exp_score = 65.0 + ((extracted_exp / min_exp) * 35.0)
        
    overall_score = (skills_score * 0.40) + (relevance_score * 0.30) + (exp_score * 0.15) + (cgpa_score * 0.15)


    
    if overall_score >= 72:
        status = 'Excellent Match'
    elif overall_score >= 60:
        status = 'Good Match'
    elif overall_score >= 45:
        status = 'Average Match'
    else:
        status = 'Low Match'
        
    candidate_name = extract_name(resume_text, filename)
    email = extract_email(resume_text)
    if not OPENROUTER_API_KEY:
        summary = "Regex match completed. Configure OPENROUTER_API_KEY in .env for AI summary."
    else:
        matched_str = ", ".join(matched_skills) if matched_skills else "None"
        missing_str = ", ".join(missing_skills) if missing_skills else "None"
        summary = f"API rate-limited. Fallback analysis: Candidate has {extracted_exp} years of experience and {extracted_cgpa if extracted_cgpa else 'N/A'} CGPA. Matched skills: {matched_str}. Missing skills: {missing_str}."
    
    return {
        'overall_score': round(overall_score, 1),
        'status': status,
        'candidate_name': candidate_name,
        'skills_score': round(skills_score, 1),
        'relevance_score': round(relevance_score, 1),
        'exp_score': round(exp_score, 1),
        'cgpa_score': round(cgpa_score, 1),
        'extracted_cgpa': str(extracted_cgpa if extracted_cgpa is not None else "Not Found"),
        'extracted_exp': extracted_exp,
        'matched_skills_str': ",".join(matched_skills),
        'missing_skills_str': ",".join(missing_skills),
        'summary': summary,
        'email': email or ''
    }


# --- Routes ---

@app.route('/analyze', methods=['POST'])
def analyze():
    # 1. Validation Checks
    if 'resume' not in request.files:
        return jsonify({'error': 'No resume file uploaded'}), 400
        
    file = request.files['resume']
    if file.filename == '':
        return jsonify({'error': 'No resume file selected'}), 400
        
    # Retrieve form parameters
    job_title = request.form.get('jobTitle', '')
    skills_raw = request.form.get('skills', '')
    min_cgpa_str = request.form.get('minCgpa', '')
    min_exp_str = request.form.get('minExperience', '')
    job_desc = request.form.get('jobDesc', '')
    
    # Safe float and int conversions
    try:
        min_cgpa = float(min_cgpa_str) if min_cgpa_str else 0.0
    except ValueError:
        min_cgpa = 0.0
        
    try:
        min_exp = int(min_exp_str) if min_exp_str else 0
    except ValueError:
        min_exp = 0
        
    # Read upload file
    filename = file.filename
    file_bytes = file.read()
    
    # 2. Basic file type check
    ext = filename.split('.')[-1].lower()
    if ext not in ['pdf', 'docx', 'txt', 'png', 'jpg', 'jpeg']:
        return jsonify({'error': 'Unsupported file format. Please upload PDF, DOCX, TXT or images (PNG, JPG, JPEG).'}), 400
        
    # 3. Screen Resume
    result = screen_resume(file_bytes, filename, job_title, skills_raw, min_cgpa, min_exp, job_desc)
        
    return jsonify({
        'overallScore': result['overall_score'],
        'status': result['status'],
        'metrics': {
            'skillsMatch': result['skills_score'],
            'projectRelevance': result['relevance_score'],
            'experienceCheck': result['exp_score'],
            'cgpaRequirement': result['cgpa_score']
        },
        'details': {
            'extractedCgpa': result['extracted_cgpa'],
            'minCgpa': min_cgpa,
            'extractedExperience': result['extracted_exp'],
            'minExperience': min_exp,
            'matchedSkills': [s.strip() for s in result['matched_skills_str'].split(',') if s.strip()] if result['matched_skills_str'] else [],
            'missingSkills': [s.strip() for s in result['missing_skills_str'].split(',') if s.strip()] if result['missing_skills_str'] else [],
            'summary': result['summary'],
            'email': result['email']
        }
    })

@app.route('/analyze_bulk', methods=['POST'])
def analyze_bulk():
    # 1. Validation Checks
    if 'resumes' not in request.files:
        return jsonify({'error': 'No resume files uploaded'}), 400
        
    files = request.files.getlist('resumes')
    if not files or files[0].filename == '':
        return jsonify({'error': 'No resume files selected'}), 400
        
    # Retrieve form parameters
    job_title = request.form.get('jobTitle', '')
    skills_raw = request.form.get('skills', '')
    min_cgpa_str = request.form.get('minCgpa', '')
    min_exp_str = request.form.get('minExperience', '')
    job_desc = request.form.get('jobDesc', '')
    
    # Safe float and int conversions
    try:
        min_cgpa = float(min_cgpa_str) if min_cgpa_str else 0.0
    except ValueError:
        min_cgpa = 0.0
        
    try:
        min_exp = int(min_exp_str) if min_exp_str else 0
    except ValueError:
        min_exp = 0
        
    # Create the job configuration record in database
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("""
        INSERT INTO jobs (title, skills, min_cgpa, min_experience, description)
        VALUES (?, ?, ?, ?, ?)
    """, (job_title, skills_raw, min_cgpa, min_exp, job_desc))
    job_id = cursor.lastrowid
    
    # Pre-read files on the main thread to prevent concurrency issues reading file objects
    file_data = []
    for file in files:
        filename = file.filename
        file_bytes = file.read()
        file_data.append((filename, file_bytes))
        
    # Define a helper function to screen a single file in a thread
    def process_single_file(item):
        filename, file_bytes = item
        ext = filename.split('.')[-1].lower()
        if ext not in ['pdf', 'docx', 'txt', 'png', 'jpg', 'jpeg']:
            return {
                'filename': filename,
                'is_supported': False,
                'res': None
            }
        try:
            res = screen_resume(file_bytes, filename, job_title, skills_raw, min_cgpa, min_exp, job_desc)
            return {
                'filename': filename,
                'is_supported': True,
                'res': res
            }
        except Exception as e:
            print(f"Error processing {filename} in thread: {str(e)}")
            return {
                'filename': filename,
                'is_supported': True,
                'error': str(e),
                'res': {
                    'overall_score': 0.0,
                    'status': 'Screening Error',
                    'candidate_name': filename,
                    'skills_score': 0.0,
                    'relevance_score': 0.0,
                    'exp_score': 0.0,
                    'cgpa_score': 0.0,
                    'extracted_cgpa': 'Not Found',
                    'extracted_exp': 0,
                    'matched_skills_str': '',
                    'missing_skills_str': skills_raw,
                    'summary': f"Error during concurrent screening: {str(e)}",
                    'email': ''
                }
            }
            
    from concurrent.futures import ThreadPoolExecutor
    
    # Process files concurrently with ThreadPoolExecutor
    # Cap maximum workers at min(number of files, 8) to avoid overloading the API / system resources
    max_workers = min(len(file_data), 8)
    if max_workers > 0:
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            results = list(executor.map(process_single_file, file_data))
    else:
        results = []
        
    # Write all candidate results to the SQLite database sequentially (thread-safe)
    for idx, item in enumerate(results):
        filename = item['filename']
        # Find file bytes
        file_bytes = None
        for fn, fb in file_data:
            if fn == filename:
                file_bytes = fb
                break

        if not item['is_supported']:
            cursor.execute("""
                INSERT INTO candidates (
                    job_id, name, filename, overall_score, status,
                    skills_match_score, project_relevance_score, experience_score, cgpa_score,
                    extracted_cgpa, extracted_experience, matched_skills, missing_skills, summary, email
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                job_id, filename, filename, 0.0, 'Unsupported Format',
                0.0, 0.0, 0.0, 0.0, 'N/A', 0, '', skills_raw, 'Unsupported file format. Use PDF, DOCX, TXT or PNG/JPG.', ''
            ))
            candidate_id = cursor.lastrowid
        else:
            res = item['res']
            # Auto-shortlist candidates scoring 72% or higher
            shortlist_status = 'Shortlisted' if res['overall_score'] >= 72.0 else 'Pending'

            cursor.execute("""
                INSERT INTO candidates (
                    job_id, name, filename, overall_score, status,
                    skills_match_score, project_relevance_score, experience_score, cgpa_score,
                    extracted_cgpa, extracted_experience, matched_skills, missing_skills, summary, shortlist_status, email
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                job_id, res['candidate_name'], filename, res['overall_score'], res['status'],
                res['skills_score'], res['relevance_score'], res['exp_score'], res['cgpa_score'],
                res['extracted_cgpa'], res['extracted_exp'], res['matched_skills_str'], res['missing_skills_str'], res['summary'],
                shortlist_status, res['email']
            ))
            candidate_id = cursor.lastrowid

        # Save to disk
        if file_bytes and candidate_id:
            UPLOAD_FOLDER = 'uploads'
            os.makedirs(UPLOAD_FOLDER, exist_ok=True)
            safe_fn = re.sub(r'[^a-zA-Z0-9_.-]', '_', filename)
            disk_filename = f"candidate_{candidate_id}_{safe_fn}"
            file_path = os.path.join(UPLOAD_FOLDER, disk_filename)
            try:
                with open(file_path, 'wb') as f:
                    f.write(file_bytes)
            except Exception as e:
                print(f"Failed to save resume file to disk: {str(e)}")
        
    conn.commit()
    conn.close()
    
    return jsonify({
        'status': 'success',
        'jobId': job_id
    })

# --- CRUD Endpoints ---

@app.route('/api/jobs', methods=['GET'])
def get_jobs_list():
    conn = get_db_connection()
    jobs = conn.execute("SELECT * FROM jobs ORDER BY created_at DESC").fetchall()
    
    job_list = []
    for job in jobs:
        candidate_count = conn.execute("SELECT COUNT(*) FROM candidates WHERE job_id = ?", (job['id'],)).fetchone()[0]
        job_list.append({
            'id': job['id'],
            'title': job['title'],
            'skills': job['skills'],
            'minCgpa': job['min_cgpa'],
            'minExperience': job['min_experience'],
            'description': job['description'],
            'createdAt': job['created_at'],
            'candidateCount': candidate_count
        })
    conn.close()
    return jsonify(job_list)

@app.route('/api/jobs/<int:job_id>', methods=['GET'])
def get_job_session(job_id):
    conn = get_db_connection()
    job = conn.execute("SELECT * FROM jobs WHERE id = ?", (job_id,)).fetchone()
    if not job:
        conn.close()
        return jsonify({'error': 'Job screening session not found'}), 404
        
    candidates_rows = conn.execute("""
        SELECT * FROM candidates 
        WHERE job_id = ? 
        ORDER BY overall_score DESC
    """, (job_id,)).fetchall()
    
    candidates = []
    for idx, row in enumerate(candidates_rows):
        candidates.append({
            'id': row['id'],
            'rank': idx + 1,
            'name': row['name'],
            'filename': row['filename'],
            'overallScore': row['overall_score'],
            'status': row['status'],
            'shortlistStatus': row['shortlist_status'] if 'shortlist_status' in row.keys() else 'Pending',
            'metrics': {
                'skillsMatch': row['skills_match_score'],
                'projectRelevance': row['project_relevance_score'],
                'experienceCheck': row['experience_score'],
                'cgpaRequirement': row['cgpa_score']
            },
            'details': {
                'extractedCgpa': row['extracted_cgpa'],
                'minCgpa': job['min_cgpa'],
                'extractedExperience': row['extracted_experience'],
                'minExperience': job['min_experience'],
                'matchedSkills': [s.strip() for s in row['matched_skills'].split(',') if s.strip()] if row['matched_skills'] else [],
                'missingSkills': [s.strip() for s in row['missing_skills'].split(',') if s.strip()] if row['missing_skills'] else [],
                'notes': row['notes'],
                'summary': row['summary'],
                'email': row['email'] if 'email' in row.keys() else ''
            }
        })
        
    conn.close()
    
    return jsonify({
        'jobId': job['id'],
        'jobTitle': job['title'],
        'skills': [s.strip() for s in job['skills'].split(',') if s.strip()] if job['skills'] else [],
        'minCgpa': job['min_cgpa'],
        'minExperience': job['min_experience'],
        'jobDesc': job['description'],
        'candidates': candidates
    })

@app.route('/api/candidates/<int:candidate_id>', methods=['PUT'])
def update_candidate(candidate_id):
    data = request.get_json()
    if not data:
        return jsonify({'error': 'No data provided'}), 400
        
    conn = get_db_connection()
    cursor = conn.cursor()
    
    update_fields = []
    params = []
    
    if 'notes' in data:
        update_fields.append("notes = ?")
        params.append(data['notes'])
    if 'shortlist_status' in data:
        update_fields.append("shortlist_status = ?")
        params.append(data['shortlist_status'])
    if 'email' in data:
        update_fields.append("email = ?")
        params.append(data['email'])
        
    if not update_fields:
        conn.close()
        return jsonify({'error': 'No fields to update'}), 400
        
    params.append(candidate_id)
    query = f"UPDATE candidates SET {', '.join(update_fields)} WHERE id = ?"
    cursor.execute(query, params)
    conn.commit()
    rows_affected = cursor.rowcount
    conn.close()
    
    if rows_affected == 0:
        return jsonify({'error': 'Candidate not found'}), 404
        
    return jsonify({'status': 'success', 'message': 'Candidate updated successfully'})


def send_email_via_smtp(recipient, candidate_name, job_title, company_name="NexusNLP"):
    import ssl
    import socket
    smtp_server = os.environ.get("SMTP_SERVER")
    smtp_port = os.environ.get("SMTP_PORT")
    smtp_username = os.environ.get("SMTP_USERNAME")
    smtp_password = os.environ.get("SMTP_PASSWORD")
    smtp_sender = os.environ.get("SMTP_SENDER")

    if not smtp_server or not smtp_username or not smtp_password:
        return False, "SMTP parameters missing in .env file (SMTP_SERVER, SMTP_USERNAME, SMTP_PASSWORD)."

    try:
        port = int(smtp_port) if smtp_port else 587
    except ValueError:
        port = 587

    sender_email = smtp_sender if smtp_sender else smtp_username

    subject = f"Shortlist Notification — {job_title} Position at {company_name}"
    
    # Professional HTML Email Template
    html_content = f"""
    <html>
        <body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333333; max-width: 600px; margin: 0 auto; padding: 20px; border: 1px solid #e0e0e0; border-radius: 8px;">
            <div style="background-color: #f97316; padding: 15px; border-radius: 6px 6px 0 0; text-align: center;">
                <h1 style="color: #ffffff; margin: 0; font-size: 24px;">{company_name} Recruitment</h1>
            </div>
            <div style="padding: 20px;">
                <p>Dear <strong>{candidate_name}</strong>,</p>
                <p>We are pleased to inform you that your resume has been reviewed and you have been <strong>shortlisted</strong> for the <strong>{job_title}</strong> position at <strong>{company_name}</strong>.</p>
                <p>Our recruitment team was highly impressed by your qualifications and experience. We will be in touch with you shortly to schedule the next steps of the interview process.</p>
                <p>If you have any questions in the meantime, please feel free to reply directly to this email.</p>
                <p>Best regards,<br>Recruitment Team<br>{company_name}</p>
            </div>
            <div style="background-color: #f5f5f5; padding: 10px; border-radius: 0 0 6px 6px; text-align: center; font-size: 12px; color: #777777; margin-top: 20px;">
                This is an automated notification from the {company_name} Smart Resume Screening System.
            </div>
        </body>
    </html>
    """

    msg = MIMEMultipart("alternative")
    msg["Subject"] = subject
    msg["From"] = f"{company_name} Recruitment <{sender_email}>"
    msg["To"] = recipient

    # Plain text version for fallback
    text_content = f"Dear {candidate_name},\n\nWe are pleased to inform you that your resume has been reviewed and you have been shortlisted for the {job_title} position at {company_name}.\n\nOur recruitment team will be in touch with you shortly to schedule the next steps of the interview process.\n\nBest regards,\nRecruitment Team\n{company_name}"
    
    msg.attach(MIMEText(text_content, "plain"))
    msg.attach(MIMEText(html_content, "html"))

    try:
        context = ssl.create_default_context()
        if port == 465:
            with smtplib.SMTP_SSL(smtp_server, port, context=context, timeout=12) as server:
                server.login(smtp_username, smtp_password)
                server.sendmail(sender_email, recipient, msg.as_string())
        else:
            with smtplib.SMTP(smtp_server, port, timeout=12) as server:
                server.ehlo()
                server.starttls(context=context)
                server.ehlo()
                server.login(smtp_username, smtp_password)
                server.sendmail(sender_email, recipient, msg.as_string())
        return True, "Email sent successfully."
    except smtplib.SMTPAuthenticationError:
        return False, "SMTP authentication failed. Verify your SMTP_USERNAME and SMTP_PASSWORD in the .env file. Note: For Gmail, you MUST use an App Password (16 letters), not your standard account password, and ensure 2-Step Verification is active on your Google account."
    except (smtplib.SMTPConnectError, ConnectionRefusedError):
        return False, f"Failed to connect to SMTP server '{smtp_server}' on port {port}. Please verify the SMTP server domain/port, and check if your local firewall or network blocks outgoing connections on this port."
    except ssl.SSLError:
        return False, f"SSL/TLS security negotiation failed. This usually occurs when the port ({port}) or connection method (SSL/TLS) doesn't match the SMTP server requirements. If using Gmail, try port 587 (STARTTLS) or port 465 (SSL)."
    except socket.timeout:
        return False, f"Connection to SMTP server '{smtp_server}' timed out. Please check your network connection or verify if the SMTP server is online."
    except Exception as e:
        return False, f"SMTP Error: {str(e)}"


@app.route('/api/smtp-status', methods=['GET'])
def get_smtp_status():
    import os
    smtp_server = os.environ.get("SMTP_SERVER")
    smtp_port = os.environ.get("SMTP_PORT")
    smtp_username = os.environ.get("SMTP_USERNAME")
    smtp_password = os.environ.get("SMTP_PASSWORD")
    smtp_sender = os.environ.get("SMTP_SENDER")
    
    configured = bool(smtp_server and smtp_username and smtp_password)
    return jsonify({
        'configured': configured,
        'server': smtp_server or '',
        'port': smtp_port or '',
        'sender': smtp_sender or smtp_username or ''
    })


@app.route('/api/candidates/<int:candidate_id>/send-email', methods=['POST'])
def send_candidate_email(candidate_id):
    conn = get_db_connection()
    candidate = conn.execute("SELECT c.*, j.title as job_title FROM candidates c JOIN jobs j ON c.job_id = j.id WHERE c.id = ?", (candidate_id,)).fetchone()
    if not candidate:
        conn.close()
        return jsonify({'error': 'Candidate not found'}), 404

    data = request.get_json() or {}
    email_address = data.get('email', '').strip()
    company_name = data.get('company', 'NexusNLP').strip() or 'NexusNLP'
    
    if not email_address:
        email_address = (candidate['email'] or '').strip()
        
    if not email_address:
        conn.close()
        return jsonify({'error': 'No email address found for this candidate. Please update the email field first.'}), 400

    # If the email address was updated/manually supplied, save it in the database
    if email_address != candidate['email']:
        conn.execute("UPDATE candidates SET email = ? WHERE id = ?", (email_address, candidate_id))
        conn.commit()

    conn.close()

    smtp_server = os.environ.get("SMTP_SERVER")
    if smtp_server:
        success, msg = send_email_via_smtp(email_address, candidate['name'], candidate['job_title'], company_name)
        if success:
            return jsonify({'status': 'success', 'message': f'Shortlist email sent to {email_address}!'})
        else:
            return jsonify({'error': f'Failed to send email via SMTP: {msg}. Please check your .env settings or backend logs.'}), 500
    else:
        # Mock Email Delivery Logging
        print("\n" + "="*50)
        print(f"MOCK EMAIL DELIVERY (SMTP NOT CONFIGURED)")
        print(f"To: {candidate['name']} <{email_address}>")
        print(f"Subject: Shortlist Notification — {candidate['job_title']} Position at {company_name}")
        print(f"Content:")
        print(f"  Dear {candidate['name']},")
        print(f"  We are pleased to inform you that you have been shortlisted for the {candidate['job_title']} position at {company_name}.")
        print("="*50 + "\n")
        
        return jsonify({
            'status': 'success',
            'message': f'Mock email logged to server console for {email_address} (SMTP not configured).',
            'mock': True
        })


@app.route('/api/jobs/<int:job_id>/notify-shortlisted', methods=['POST'])
def notify_all_shortlisted(job_id):
    data = request.get_json() or {}
    company_name = data.get('company', 'NexusNLP').strip() or 'NexusNLP'
    
    # 1. Fetch job and all shortlisted candidates
    conn = get_db_connection()
    job = conn.execute("SELECT * FROM jobs WHERE id = ?", (job_id,)).fetchone()
    if not job:
        conn.close()
        return jsonify({'error': 'Job session not found'}), 404
        
    candidates = conn.execute("""
        SELECT * FROM candidates 
        WHERE job_id = ? AND shortlist_status = 'Shortlisted'
    """, (job_id,)).fetchall()
    conn.close()
    
    if not candidates:
        return jsonify({'status': 'success', 'message': 'No shortlisted candidates found to notify.', 'sent_count': 0})
        
    # Check if SMTP is configured
    smtp_server = os.environ.get("SMTP_SERVER")
    
    # Helper to notify a single candidate
    def notify_candidate(cand):
        email_address = (cand['email'] or '').strip()
        if not email_address:
            return {'id': cand['id'], 'name': cand['name'], 'email': '', 'success': False, 'error': 'No email address found.'}
            
        if smtp_server:
            success, err_msg = send_email_via_smtp(email_address, cand['name'], job['title'], company_name)
            return {
                'id': cand['id'],
                'name': cand['name'],
                'email': email_address,
                'success': success,
                'error': None if success else err_msg,
                'mocked': False
            }
        else:
            # Mock Logging
            print("\n" + "="*50)
            print(f"MOCK BULK EMAIL DELIVERY (SMTP NOT CONFIGURED)")
            print(f"To: {cand['name']} <{email_address}>")
            print(f"Subject: Shortlist Notification — {job['title']} Position at {company_name}")
            print("="*50 + "\n")
            return {
                'id': cand['id'],
                'name': cand['name'],
                'email': email_address,
                'success': True,
                'error': 'Mock email logged (SMTP not configured).',
                'mocked': True
            }
            
    # Send emails concurrently using ThreadPoolExecutor
    from concurrent.futures import ThreadPoolExecutor
    max_workers = min(len(candidates), 10)
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        results = list(executor.map(notify_candidate, candidates))
        
    sent_count = sum(1 for r in results if r['success'])
    failed_count = len(results) - sent_count
    
    return jsonify({
        'status': 'success',
        'message': f'Notified {sent_count} candidate(s) successfully. {failed_count} failed.',
        'sent_count': sent_count,
        'failed_count': failed_count,
        'details': results
    })


@app.route('/api/candidates/<int:candidate_id>', methods=['DELETE'])
def delete_candidate(candidate_id):
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM candidates WHERE id = ?", (candidate_id,))
    conn.commit()
    rows_affected = cursor.rowcount
    conn.close()
    
    if rows_affected == 0:
        return jsonify({'error': 'Candidate not found'}), 404
        
    return jsonify({'status': 'success', 'message': 'Candidate record deleted'})

@app.route('/api/jobs/<int:job_id>', methods=['DELETE'])
def delete_job_session(job_id):
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("PRAGMA foreign_keys = ON;")
    cursor.execute("DELETE FROM jobs WHERE id = ?", (job_id,))
    conn.commit()
    rows_affected = cursor.rowcount
    conn.close()
    
    if rows_affected == 0:
        return jsonify({'error': 'Job screening session not found'}), 404
        
    return jsonify({'status': 'success', 'message': 'Job session and candidates deleted'})

@app.route('/api/candidates/<int:candidate_id>/resume', methods=['GET'])
def get_candidate_resume(candidate_id):
    from flask import send_file
    conn = get_db_connection()
    row = conn.execute("SELECT filename FROM candidates WHERE id = ?", (candidate_id,)).fetchone()
    conn.close()
    if not row:
        return jsonify({'error': 'Candidate not found'}), 404
    filename = row['filename']
    safe_fn = re.sub(r'[^a-zA-Z0-9_.-]', '_', filename)
    file_path = os.path.join('uploads', f"candidate_{candidate_id}_{safe_fn}")
    if not os.path.exists(file_path):
        return jsonify({'error': 'Resume file not found on server'}), 404
        
    ext = filename.split('.')[-1].lower()
    mimetype = 'application/octet-stream'
    if ext == 'pdf':
        mimetype = 'application/pdf'
    elif ext in ['png', 'jpg', 'jpeg']:
        mimetype = f'image/{ext if ext != "jpg" else "jpeg"}'
    elif ext == 'docx':
        mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    elif ext == 'txt':
        mimetype = 'text/plain'
        
    return send_file(file_path, mimetype=mimetype, as_attachment=False, download_name=filename)

@app.route('/api/candidates/<int:candidate_id>/chart.png', methods=['GET'])
def get_candidate_chart(candidate_id):
    # pyrefly: ignore [missing-import]
    import matplotlib
    matplotlib.use('Agg')
    # pyrefly: ignore [missing-import]
    import matplotlib.pyplot as plt
    import io
    from flask import send_file

    conn = get_db_connection()
    candidate = conn.execute("""
        SELECT skills_match_score, project_relevance_score, experience_score, cgpa_score, name 
        FROM candidates WHERE id = ?
    """, (candidate_id,)).fetchone()
    conn.close()

    if not candidate:
        return jsonify({'error': 'Candidate not found'}), 404

    categories = ['Skills\nMatch', 'Project\nRelevance', 'Experience\nCheck', 'CGPA\nReq']
    categories = categories[::-1]
    
    scores = [
        candidate['skills_match_score'],
        candidate['project_relevance_score'],
        candidate['experience_score'],
        candidate['cgpa_score']
    ]
    scores = scores[::-1]

    fig, ax = plt.subplots(figsize=(6, 2.8), dpi=120)
    
    fig.patch.set_facecolor('none')
    ax.patch.set_facecolor('none')

    colors = ['#ffedd5', '#ffbe98', '#ff9453', '#f97316']
    
    bars = ax.barh(categories, scores, color=colors, height=0.55, edgecolor='none')

    ax.set_xlim(0, 105)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['bottom'].set_visible(False)
    ax.spines['left'].set_visible(False)
    ax.xaxis.set_visible(False)

    ax.grid(False)

    theme = request.args.get('theme', 'dark')
    if theme == 'light':
        label_color = '#44403c'
        text_color = '#292524'
    else:
        label_color = '#a8a29e'
        text_color = '#f5f5f4'

    ax.tick_params(axis='y', colors=label_color, labelsize=9, length=0)

    for bar in bars:
        width = bar.get_width()
        ax.text(
            width + 2,
            bar.get_y() + bar.get_height() / 2,
            f'{int(width)}%',
            ha='left',
            va='center',
            color=text_color,
            fontweight='bold',
            fontsize=9
        )

    plt.tight_layout()

    buf = io.BytesIO()
    plt.savefig(buf, format='png', facecolor=fig.get_facecolor(), edgecolor='none', bbox_inches='tight', pad_inches=0.1)
    buf.seek(0)
    plt.close(fig)

    return send_file(buf, mimetype='image/png')

# --- Authentication Endpoints ---

@app.route('/api/login', methods=['POST'])
def login():
    data = request.get_json()
    email = data.get('email', '').strip().lower()
    password = data.get('password', '')
    
    if not email or not password:
        return jsonify({'error': 'Email and password are required'}), 400
    
    conn = get_db_connection()
    user = conn.execute("SELECT * FROM users WHERE email = ?", (email,)).fetchone()
    conn.close()
    
    if not user:
        return jsonify({'error': 'Invalid email or password'}), 401
    
    if user['password_hash'] != hash_password(password):
        return jsonify({'error': 'Invalid email or password'}), 401
    
    return jsonify({
        'status': 'success',
        'user': {
            'id': user['id'],
            'fullName': user['full_name'],
            'email': user['email']
        }
    })

@app.route('/api/register', methods=['POST'])
def register():
    data = request.get_json()
    full_name = data.get('fullName', '').strip()
    email = data.get('email', '').strip().lower()
    password = data.get('password', '')
    
    if not full_name or not email or not password:
        return jsonify({'error': 'Full name, email, and password are required'}), 400
    
    if len(password) < 6:
        return jsonify({'error': 'Password must be at least 6 characters'}), 400
    
    conn = get_db_connection()
    existing = conn.execute("SELECT id FROM users WHERE email = ?", (email,)).fetchone()
    if existing:
        conn.close()
        return jsonify({'error': 'An account with this email already exists'}), 409
    
    cursor = conn.cursor()
    cursor.execute(
        "INSERT INTO users (full_name, email, password_hash) VALUES (?, ?, ?)",
        (full_name, email, hash_password(password))
    )
    conn.commit()
    user_id = cursor.lastrowid
    conn.close()
    
    return jsonify({
        'status': 'success',
        'user': {
            'id': user_id,
            'fullName': full_name,
            'email': email
        }
    })

if __name__ == "__main__":
    init_db()
    port = int(os.environ.get("PORT", 5000))
    print(f"Starting smart resume screening backend on http://0.0.0.0:{port}...")
    app.run(host="0.0.0.0", port=port, debug=False)
